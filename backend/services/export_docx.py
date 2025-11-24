from docx import Document

def export_docx(project, sections):
    doc = Document()
    doc.add_heading(project.title, level=0)

    for sec in sections:
        doc.add_heading(sec.title, level=1)
        doc.add_paragraph(sec.content or "")

    path = f"exports/{project.title.replace(' ', '_')}.docx"
    doc.save(path)
    return path
